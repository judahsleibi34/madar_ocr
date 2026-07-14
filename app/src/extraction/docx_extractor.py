import docx


def extract_docx_text(docx_path: str) -> list:
    doc = docx.Document(docx_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append({"text": text})

    return paragraphs


def extract_docx_tables(docx_path: str) -> list:
    doc = docx.Document(docx_path)
    tables_data = []

    for table_index, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables_data.append({
            "table_index": table_index,
            "rows": rows
        })

    return tables_data
